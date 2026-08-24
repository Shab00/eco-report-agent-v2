from __future__ import annotations

import os
import re
import unicodedata
import uuid
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from app.models import AuthorInfo, PEAReport, SiteInfo, SurveyConditions

DARK_GREEN = "2D5A27"
LIGHT_GREEN = "E8F5E9"
SAGE_GREEN = "C8D8B0"
RULE_GREEN = "90C226"  # sampled from templates/template1.pdf & template2.pdf rule lines
WHITE = "FFFFFF"

DISCLAIMER_TEMPLATE = (
    "This report dated {date} has been prepared for {client} (the "
    "'Client') in accordance with the terms and conditions of "
    "appointment (the 'Appointment') between the Client and "
    "{author} (Greenlight Ecology Ltd) for the purposes specified "
    "in the Appointment. For avoidance of doubt, no other person(s) "
    "may use or rely upon this report or its contents, and Greenlight "
    "Ecology Ltd accepts no responsibility for any such use or "
    "reliance thereon by any other third party."
)

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "generated")
LOGO_PATH_ENV = os.environ.get("LOGO_PATH")


# --------------------------------------------------------------------------
# OXML helpers
# --------------------------------------------------------------------------


# CT_TcPr requires its children in this order (ECMA-376); appending blindly
# can put w:shd before w:tcBorders and produce a .docx Word refuses to open.
_TCPR_TAG_SEQ = (
    "w:cnfStyle",
    "w:tcW",
    "w:gridSpan",
    "w:hMerge",
    "w:vMerge",
    "w:tcBorders",
    "w:shd",
    "w:noWrap",
    "w:tcMar",
    "w:textDirection",
    "w:tcFitText",
    "w:vAlign",
    "w:hideMark",
    "w:headers",
    "w:cellIns",
    "w:cellDel",
    "w:cellMerge",
    "w:tcPrChange",
)


def _insert_tcPr_child(tcPr, element, tag: str) -> None:
    successors = _TCPR_TAG_SEQ[_TCPR_TAG_SEQ.index(tag) + 1 :]
    tcPr.insert_element_before(element, *successors)


def set_cell_background(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _insert_tcPr_child(tcPr, shd, "w:shd")


def set_cell_borders_none(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    _insert_tcPr_child(tcPr, borders, "w:tcBorders")


# CT_TblPr also requires its children in schema order.
_TBLPR_TAG_SEQ = (
    "w:tblStyle",
    "w:tblpPr",
    "w:tblOverlap",
    "w:bidiVisual",
    "w:tblStyleRowBandSize",
    "w:tblStyleColBandSize",
    "w:tblW",
    "w:jc",
    "w:tblCellSpacing",
    "w:tblInd",
    "w:tblBorders",
    "w:shd",
    "w:tblLayout",
    "w:tblCellMar",
    "w:tblLook",
    "w:tblCaption",
    "w:tblDescription",
    "w:tblPrChange",
)


def _insert_tblPr_child(tblPr, element, tag: str) -> None:
    successors = _TBLPR_TAG_SEQ[_TBLPR_TAG_SEQ.index(tag) + 1 :]
    tblPr.insert_element_before(element, *successors)


def set_table_top_rule(table, hex_color: str, size_eighths_pt: int = 12) -> None:
    """Single rule line above the table only; no other table borders."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_eighths_pt))
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), hex_color)
    borders.append(top)
    for edge in ("left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    _insert_tblPr_child(tblPr, borders, "w:tblBorders")


def add_page_number_field(paragraph, instruction: str = "PAGE") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_repeat_header_row(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_exact_row_height(row, twips: int) -> None:
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


# --------------------------------------------------------------------------
# GPT text sanitisation
# --------------------------------------------------------------------------

# XML 1.0 forbids these outright (not even valid as character references),
# so any that reach python-docx's run/text setters raise ValueError before
# a file is ever produced. Stripped here anyway as defense in depth in case
# GPT output ever contains them, rather than letting a request 500.
_ILLEGAL_XML_CHARS_RE = re.compile(
    "[\x00-\x08\x0b\x0c\x0e-\x1f\x7f￾￿]"
)

# Characters that are valid XML but that GPT commonly emits and that we
# don't want rendered literally in a formal report (smart punctuation,
# markdown emphasis markers, bullet glyphs).
_CHAR_REPLACEMENTS = {
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "–": "-",
    "—": "-",
    "…": "...",
    "•": "-",
    " ": " ",
}

_MARKDOWN_EMPHASIS_RE = re.compile(r"\*{1,3}|_{2,3}")
_MARKDOWN_HEADING_RE = re.compile(r"^#{1,6}\s*", re.MULTILINE)


def sanitize_text(text: str | None) -> str:
    """Clean freeform (typically GPT-generated) text before it is written
    into a run. Normalises unicode, swaps smart punctuation/bullets for
    plain equivalents, strips markdown emphasis/heading markers, and
    removes characters XML does not allow as content."""
    if not text:
        return ""
    cleaned = unicodedata.normalize("NFKC", text)
    for bad, good in _CHAR_REPLACEMENTS.items():
        cleaned = cleaned.replace(bad, good)
    cleaned = _MARKDOWN_HEADING_RE.sub("", cleaned)
    cleaned = _MARKDOWN_EMPHASIS_RE.sub("", cleaned)
    cleaned = _ILLEGAL_XML_CHARS_RE.sub("", cleaned)
    return cleaned


def add_text_as_paragraphs(cell, lines: list[str]) -> None:
    """Write each item in `lines` as its own paragraph, first sanitising it
    and splitting any embedded newlines (GPT sometimes returns a single
    field as a "\\n"-joined list) into further paragraphs of their own,
    rather than embedding a raw newline character inside one run."""
    cell.text = ""
    paragraph_texts: list[str] = []
    for line in lines:
        if not line:
            continue
        for sub_line in sanitize_text(line).splitlines():
            sub_line = sub_line.strip(" \t-•")
            if sub_line:
                paragraph_texts.append(sub_line)

    if not paragraph_texts:
        cell.paragraphs[0].add_run("")
        return

    for idx, paragraph_text in enumerate(paragraph_texts):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.add_run(paragraph_text)


# --------------------------------------------------------------------------
# Base document / style setup
# --------------------------------------------------------------------------


def configure_base_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Arial")


def configure_section(section, margins_cm: float | None) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    if margins_cm is not None:
        section.top_margin = Cm(margins_cm)
        section.bottom_margin = Cm(margins_cm)
        section.left_margin = Cm(margins_cm)
        section.right_margin = Cm(margins_cm)
    else:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)


def populate_body_header(section, site_name: str, logo_path: str | None = None) -> None:
    # A borderless 1x2 table gives reliable independent left/right placement
    # for text + an inline image; a right-aligned tab stop is not reliably
    # honoured by Word for inline drawings and can render the image
    # mid-line instead of flush right.
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""

    usable_width = section.page_width - section.left_margin - section.right_margin
    logo_col_width = Cm(2.5)
    text_col_width = usable_width - logo_col_width

    table = header.add_table(rows=1, cols=2, width=usable_width)
    table.autofit = False
    table.columns[0].width = text_col_width
    table.columns[1].width = logo_col_width

    text_cell = table.cell(0, 0)
    text_cell.width = text_col_width
    set_cell_borders_none(text_cell)
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    text_paragraph = text_cell.paragraphs[0]
    text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    site_run = text_paragraph.add_run(f"{sanitize_text(site_name)} PEA")
    site_run.font.size = Pt(9)
    site_run.font.color.rgb = RGBColor.from_string(DARK_GREEN)

    logo_cell = table.cell(0, 1)
    logo_cell.width = logo_col_width
    set_cell_borders_none(logo_cell)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    effective_logo_path = logo_path or LOGO_PATH_ENV
    if effective_logo_path and os.path.exists(effective_logo_path):
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(effective_logo_path, width=Cm(1.5))

    # Drop the empty paragraph Word seeds every header with, now that the
    # table is the header's only content, so the table sits flush at top.
    stray_paragraph = header.paragraphs[0]
    stray_paragraph._p.getparent().remove(stray_paragraph._p)


def populate_body_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("Greenlight Ecology Ltd")
    run.font.size = Pt(8)

    tab_stops = paragraph.paragraph_format.tab_stops
    usable_width = section.page_width - section.left_margin - section.right_margin
    tab_stops.add_tab_stop(usable_width, alignment=2)  # right-aligned tab
    center_run = paragraph.add_run("\t")
    center_run = paragraph.add_run("GPage | ")
    center_run.font.size = Pt(8)
    add_page_number_field(paragraph, instruction="PAGE \\* ROMAN")


# --------------------------------------------------------------------------
# Page builders
# --------------------------------------------------------------------------


def build_cover_page(doc: Document, site_info: SiteInfo, author_info: AuthorInfo, logo_path: str | None = None) -> None:
    section = doc.sections[0]
    configure_section(section, margins_cm=None)

    page_height_twips = section.page_height.twips
    author_row_twips = 1500
    main_row_twips = page_height_twips - author_row_twips

    # Two full-width, borderless, sage-green rows so the background reads as
    # one continuous full-bleed field: a vertically-centred main row for the
    # logo/titles, and a short bottom row anchored to the page foot for the
    # author line.
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    table.columns[0].width = section.page_width

    main_cell = table.cell(0, 0)
    main_cell.width = section.page_width
    set_cell_background(main_cell, SAGE_GREEN)
    set_cell_borders_none(main_cell)
    main_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_exact_row_height(table.rows[0], main_row_twips)

    author_cell = table.cell(1, 0)
    author_cell.width = section.page_width
    set_cell_background(author_cell, SAGE_GREEN)
    set_cell_borders_none(author_cell)
    author_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    set_exact_row_height(table.rows[1], author_row_twips)

    main_cell.paragraphs[0].text = ""

    effective_logo_path = logo_path or LOGO_PATH_ENV
    if effective_logo_path and os.path.exists(effective_logo_path):
        logo_paragraph = main_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_paragraph.add_run()
        run.add_picture(effective_logo_path, width=Cm(6))

    def add_cover_paragraph(
        text: str,
        size: int,
        bold: bool = False,
        italic: bool = False,
        color: str | None = None,
        font_name: str | None = None,
        space_before: int = 6,
    ):
        p = main_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        if font_name:
            r.font.name = font_name
        return p

    add_cover_paragraph("GREENLIGHT ECOLOGY LTD", 16, bold=True, color=DARK_GREEN, space_before=12)
    add_cover_paragraph(
        "Preliminary Ecological Appraisal", 22, italic=True, font_name="Times New Roman", space_before=18
    )
    add_cover_paragraph(site_info.site_name, 13, color=DARK_GREEN, space_before=12)

    author_cell.text = ""
    author_paragraph = author_cell.paragraphs[0]
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    author_paragraph.paragraph_format.left_indent = Cm(1.5)
    author_run = author_paragraph.add_run(
        f"{author_info.name} {author_info.credentials}, {author_info.role}"
    )
    author_run.font.size = Pt(11)


def build_version_control_page(doc: Document, site_info: SiteInfo, author_info: AuthorInfo) -> None:
    today = datetime.now().strftime("%d/%m/%Y")

    heading = doc.add_paragraph()
    run = heading.add_run("Preliminary Ecological Appraisal")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(DARK_GREEN)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(site_info.site_name)
    sub_run.font.size = Pt(14)
    sub_run.italic = True

    doc.add_paragraph()

    def add_meta_line(label: str, value: str):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(value)

    add_meta_line("Prepared By", f"{author_info.name}, Greenlight Ecology Ltd")
    add_meta_line("Prepared For", site_info.prepared_for)
    add_meta_line("Doc Ref", site_info.doc_ref)
    add_meta_line("Author", f"{author_info.name} {author_info.credentials}, {author_info.role}")

    doc.add_paragraph()

    version_heading = doc.add_paragraph()
    vh_run = version_heading.add_run("Version Control")
    vh_run.bold = True
    vh_run.font.size = Pt(12)

    columns = ["Version", "Date", "Author", "Checker", "Approver", "Changes"]
    table = doc.add_table(rows=2, cols=len(columns))
    set_table_top_rule(table, RULE_GREEN)

    header_row = table.rows[0]
    for idx, label in enumerate(columns):
        cell = header_row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_background(cell, DARK_GREEN)
    set_repeat_header_row(header_row)

    data_row = table.rows[1]
    values = ["P01", today, author_info.name, "", "", "First issue"]
    for idx, value in enumerate(values):
        data_row.cells[idx].text = value

    doc.add_paragraph()

    disclaimer_text = DISCLAIMER_TEMPLATE.format(
        date=today, client=site_info.prepared_for, author=author_info.name
    )
    disclaimer_p = doc.add_paragraph()
    disclaimer_run = disclaimer_p.add_run(disclaimer_text)
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.italic = True

    doc.add_page_break()


def build_survey_conditions(doc: Document, survey_conditions: SurveyConditions, author_info: AuthorInfo) -> None:
    intro = doc.add_paragraph()
    intro.add_run(
        f"The site survey was undertaken by {author_info.name} "
        f"{author_info.credentials} Consultant Ecologist."
    )

    doc.add_paragraph()

    columns = [
        "Date of survey",
        "Temperature (°C)",
        "Humidity (%)",
        "Cloud Cover (%)",
        "Wind (km/h)",
        "Rain",
    ]
    table = doc.add_table(rows=2, cols=len(columns))
    table.style = "Table Grid"

    header_row = table.rows[0]
    for idx, label in enumerate(columns):
        cell = header_row.cells[idx]
        cell.text = ""
        r = cell.paragraphs[0].add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_background(cell, DARK_GREEN)
    set_repeat_header_row(header_row)

    data_row = table.rows[1]
    values = [
        survey_conditions.date,
        str(survey_conditions.temperature_c),
        str(survey_conditions.humidity_percent),
        str(survey_conditions.cloud_cover_percent),
        str(survey_conditions.wind_kmh),
        survey_conditions.rain,
    ]
    for idx, value in enumerate(values):
        data_row.cells[idx].text = value

    doc.add_paragraph()


TOPIC_SPECS = [
    "Habitats and Plants",
    "Locality and Designated Sites",
    "Invasive / Non-native species",
    "Invertebrates",
    "Bats",
    "Birds",
    "Reptiles",
    "Amphibians",
    "Badger",
    "Riparian Animals",
    "Hazel Dormouse",
    "Other (e.g. Hedgehog)",
]


def _add_band_row(table, label: str) -> None:
    row = table.add_row()
    merged = row.cells[0].merge(row.cells[1])
    merged.text = ""
    p = merged.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    set_cell_background(merged, DARK_GREEN)


_row_toggle = {"n": 0}


def _add_content_row(table, label: str, lines: list[str]) -> None:
    row = table.add_row()
    label_cell, content_cell = row.cells[0], row.cells[1]

    label_cell.text = ""
    lp = label_cell.paragraphs[0]
    lr = lp.add_run(sanitize_text(label))
    lr.italic = True
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_text_as_paragraphs(content_cell, lines)

    shade = LIGHT_GREEN if _row_toggle["n"] % 2 == 0 else WHITE
    _row_toggle["n"] += 1
    set_cell_background(label_cell, shade)
    set_cell_background(content_cell, shade)


def _species_summary_lines(section) -> list[str]:
    return [
        f"EPSL / MAGIC Data: {section.epsl_data}",
        f"Habitat Suitability: {section.habitat_suitability}",
        f"Survey Findings: {section.survey_findings}",
    ]


def _species_recommendations_lines(section) -> list[str]:
    lines = [section.recommendations]
    if section.biodiversity_enhancements:
        lines.append(f"Suggested biodiversity enhancements: {section.biodiversity_enhancements}")
    return lines


def build_assessment_table(doc: PEAReport, report: PEAReport) -> None:
    _row_toggle["n"] = 0
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True

    header_row = table.rows[0]
    headers = [
        "Ecological Survey Factor",
        "Conclusion, Impact or Recommendations",
    ]
    for idx, label in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = ""
        r = cell.paragraphs[0].add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_background(cell, DARK_GREEN)
    set_repeat_header_row(header_row)

    # 1. Habitats and Plants
    _add_band_row(table, "Habitats and Plants")
    habitats = report.habitats
    summary_lines = [habitats.site_context] + list(habitats.ukhab_descriptions)
    _add_content_row(table, "Summary of Survey Findings", summary_lines)
    _add_content_row(table, "Foreseen Impacts", [habitats.foreseen_impacts])
    _add_content_row(table, "Recommendations", [habitats.recommendations])

    # 2. Locality and Designated Sites
    _add_band_row(table, "Locality and Designated Sites")
    des = report.designated_sites
    des_summary = [
        f"On-site designations: {des.on_site_designations}",
        f"Statutory sites within 2km: {des.statutory_sites}",
        f"Non-statutory designated sites: {des.non_statutory_sites}",
    ]
    _add_content_row(table, "Summary of Survey Findings", des_summary)
    _add_content_row(table, "Foreseen Impacts", [des.foreseen_impacts])
    _add_content_row(table, "Recommendations", [des.recommendations])

    # 3-12. Species sections
    species_map = [
        ("Invasive / Non-native species", report.invasive_species),
        ("Invertebrates", report.invertebrates),
        ("Bats", report.bats),
        ("Birds", report.birds),
        ("Reptiles", report.reptiles),
        ("Amphibians", report.amphibians),
        ("Badger", report.badger),
        ("Riparian Animals", report.riparian),
        ("Hazel Dormouse", report.dormouse),
        ("Other (e.g. Hedgehog)", report.hedgehog),
    ]
    for label, section in species_map:
        _add_band_row(table, label)
        _add_content_row(table, "Summary of Survey Findings", _species_summary_lines(section))
        _add_content_row(table, "Foreseen Impacts", [section.foreseen_impacts])
        _add_content_row(table, "Recommendations", _species_recommendations_lines(section))


def build_appendix_placeholders(doc: Document) -> None:
    def add_placeholder_page(title: str, body_text: str):
        doc.add_page_break()
        heading = doc.add_paragraph()
        r = heading.add_run(title)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor.from_string(DARK_GREEN)

        box = doc.add_table(rows=1, cols=1)
        box.style = "Table Grid"
        cell = box.cell(0, 0)
        cell.text = ""
        set_cell_background(cell, LIGHT_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(body_text)
        run.italic = True

    add_placeholder_page("Appendix 1: Survey/Habitat Map", "[Survey/Habitat map to be inserted]")
    add_placeholder_page("Appendix 2: Location Map", "[Location map to be inserted]")
    add_placeholder_page("Appendix 3: Proposed Plan", "[Proposed plan to be inserted]")

    doc.add_page_break()
    heading = doc.add_paragraph()
    r = heading.add_run("Appendix 4: Photos")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(DARK_GREEN)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_row = table.rows[0]
    for idx, label in enumerate(["Photograph", "Description"]):
        cell = header_row.cells[idx]
        cell.text = ""
        r = cell.paragraphs[0].add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_background(cell, DARK_GREEN)


# --------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------


def build_document(report: PEAReport, logo_path: str | None = None) -> str:
    doc = Document()
    configure_base_styles(doc)

    build_cover_page(doc, report.site_info, report.author_info, logo_path)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, margins_cm=2.54)
    populate_body_header(body_section, report.site_info.site_name, logo_path)
    populate_body_footer(body_section)

    build_version_control_page(doc, report.site_info, report.author_info)
    build_survey_conditions(doc, report.survey_conditions, report.author_info)
    build_assessment_table(doc, report)
    build_appendix_placeholders(doc)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filename = f"{report.report_id}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    return output_path
